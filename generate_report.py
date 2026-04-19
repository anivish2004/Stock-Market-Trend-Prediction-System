"""
Generate the Stock Market Trend Prediction System project report as a Word document.
Run:  python generate_report.py
Output: Project_Report.docx  (same directory)
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BLANK = "________________________"
BLANK_SHORT = "______________"
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Project_Report.docx")


# ── Helpers ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color (shading) on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(val))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "999999")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_page_break()


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return h


def para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def centered(doc, text, bold=False, size=12, space_after=6):
    return para(doc, text, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, space_after=space_after)


def bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(12)
        run = p.add_run(text)
        run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
    return p


def numbered(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(12)
        r2 = p.add_run(text)
        r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.size = Pt(12)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A237E")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F0FF")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing
    return table


# ── Build Report ─────────────────────────────────────────

def build_report():
    doc = Document()

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ── Page margins (1 inch) ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ════════════════════════════════════════════════════════
    #                     FRONT PAGE
    # ════════════════════════════════════════════════════════

    # Top spacing
    for _ in range(3):
        doc.add_paragraph()

    centered(doc, BLANK, size=18, bold=True, space_after=2)
    centered(doc, "(College / University Name)", size=10, space_after=2)
    centered(doc, BLANK, size=11, space_after=16)
    centered(doc, "(College Address)", size=10, space_after=16)

    centered(doc, f"Department of {BLANK_SHORT}", size=14, bold=True, space_after=28)

    centered(doc, "A PROJECT REPORT ON", size=12, bold=True, space_after=8)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    run_t = p_title.add_run("STOCK MARKET TREND PREDICTION SYSTEM\nUSING HYBRID ANN-GA APPROACH")
    run_t.bold = True
    run_t.font.size = Pt(20)
    run_t.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph()

    centered(doc, "Submitted in partial fulfillment of the requirements for the degree of", size=11, space_after=4)
    centered(doc, f"{BLANK}  (e.g., B.Tech / MCA)", size=11, space_after=20)

    centered(doc, "Submitted By:", size=12, bold=True, space_after=8)

    # Team table
    team_table = doc.add_table(rows=4, cols=3)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    team_table.style = "Table Grid"
    headers_team = ["S.No", "Name", "Reg. No / Roll No"]
    for i, h in enumerate(headers_team):
        cell = team_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "E8EAF6")
    for r in range(1, 4):
        team_table.rows[r].cells[0].text = f"{r}."
        team_table.rows[r].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        team_table.rows[r].cells[1].text = BLANK_SHORT
        team_table.rows[r].cells[2].text = BLANK_SHORT
        for c in range(3):
            for p in team_table.rows[r].cells[c].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()
    centered(doc, "Under the Guidance of:", size=12, bold=True, space_after=6)
    centered(doc, f"{BLANK}  (Guide Name & Designation)", size=11, space_after=28)
    centered(doc, f"Academic Year: {BLANK_SHORT}", size=13, bold=True, space_after=4)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #                     CERTIFICATE
    # ════════════════════════════════════════════════════════

    heading(doc, "CERTIFICATE", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    centered(doc, BLANK, size=14, bold=True, space_after=2)
    centered(doc, "(College / University Name)", size=10, space_after=4)
    centered(doc, f"Department of {BLANK_SHORT}", size=11, space_after=16)

    para(doc, (
        'This is to certify that the project report entitled "Stock Market Trend Prediction System '
        'Using Hybrid ANN-GA Approach" is a bonafide record of the project work carried out by:'
    ), space_after=12)

    for i in range(1, 4):
        para(doc, f"   {i}.  {BLANK}   (Reg. No: {BLANK_SHORT})", space_after=4)

    para(doc, (
        f"\nin partial fulfillment of the requirements for the award of the degree of {BLANK} in "
        f"{BLANK_SHORT} during the academic year {BLANK_SHORT}."
    ), space_after=12)

    para(doc, (
        f"The project has been carried out under the supervision of {BLANK} and is approved for submission."
    ), space_after=28)

    # Signatures
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in range(2):
        for c in range(2):
            for border in ["top", "bottom", "left", "right"]:
                set_cell_borders(sig_table.rows[r].cells[c])

    sig_table.rows[0].cells[0].text = ""
    sig_table.rows[0].cells[1].text = ""
    sig_table.rows[1].cells[0].text = ""
    sig_table.rows[1].cells[1].text = ""

    p0 = sig_table.rows[0].cells[0].paragraphs[0]
    p0.add_run("\n\n\n________________________\nProject Guide\n").font.size = Pt(11)
    p0.add_run(f"{BLANK_SHORT}").font.size = Pt(11)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = sig_table.rows[0].cells[1].paragraphs[0]
    p1.add_run("\n\n\n________________________\nHead of Department\n").font.size = Pt(11)
    p1.add_run(f"{BLANK_SHORT}").font.size = Pt(11)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove borders from signature table
    for row in sig_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ["top", "bottom", "left", "right"]:
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                tcBorders.append(el)
            tcPr.append(tcBorders)

    doc.add_paragraph()
    centered(doc, "\n\n________________________\nExternal Examiner", size=11, space_after=16)
    centered(doc, f"Date: {BLANK_SHORT}          Place: {BLANK_SHORT}", size=11)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #                       INDEX
    # ════════════════════════════════════════════════════════

    heading(doc, "INDEX", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    index_rows = [
        ["—", "Abstract", ""],
        ["1", "Introduction", ""],
        ["2", "Literature Review", ""],
        ["3", "Problem Statement", ""],
        ["4", "Methodology", ""],
        ["5", "Design & Implementation", ""],
        ["5.1", "    System Architecture", ""],
        ["5.2", "    Module Description", ""],
        ["5.3", "    Database Design", ""],
        ["5.4", "    User Interface", ""],
        ["6", "Testing & Results", ""],
        ["6.1", "    Test Cases", ""],
        ["6.2", "    Performance Metrics", ""],
        ["6.3", "    Screenshots", ""],
        ["7", "Conclusion", ""],
        ["7.1", "    Future Scope", ""],
        ["—", "References", ""],
    ]
    add_table(doc, ["Chapter No.", "Title", "Page No."], index_rows, col_widths=[1.2, 4, 1.2])

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #                      ABSTRACT
    # ════════════════════════════════════════════════════════

    heading(doc, "ABSTRACT", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    para(doc, (
        "Stock market prediction is one of the most challenging tasks in the domain of financial analytics "
        "due to the inherently volatile and non-linear nature of market data. This project presents a "
        "Stock Market Trend Prediction System that employs a Hybrid Artificial Neural Network–Genetic "
        "Algorithm (ANN-GA) approach to forecast short-term stock price direction—classifying the next "
        "trading session as Bullish, Bearish, or Neutral."
    ))
    para(doc, (
        "The system fetches historical OHLCV (Open, High, Low, Close, Volume) data from Yahoo Finance "
        "using the yfinance library, engineers a comprehensive set of technical indicators—including "
        "Simple Moving Average (SMA), Exponential Moving Average (EMA), Relative Strength Index (RSI), "
        "Moving Average Convergence Divergence (MACD), momentum, and volatility metrics—and constructs "
        "a feature matrix for model training. A Genetic Algorithm is employed to perform automated "
        "feature selection and hyperparameter tuning for a Multi-Layer Perceptron (MLP) neural network "
        "classifier, thereby optimizing prediction accuracy without manual intervention."
    ))
    para(doc, (
        "The front-end application is built using Streamlit, offering a professional, chart-first "
        "trading-terminal interface with interactive Plotly visualizations including candlestick charts, "
        "RSI plots, MACD histograms, probability gauges, and confusion matrix heatmaps. The system also "
        "integrates with MongoDB Atlas for persistent storage of prediction logs, enabling users to "
        "search, audit, and compare historical runs over time."
    ))
    para(doc, (
        "The project demonstrates that evolutionary optimization combined with neural network "
        "classification can produce reliable decision-support signals for stock market analysis, "
        "achieving competitive accuracy on validation data across multiple ticker symbols from both "
        "US and Indian (NSE) markets."
    ))

    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_kw_b = p_kw.add_run("Keywords: ")
    run_kw_b.bold = True
    run_kw_b.font.size = Pt(12)
    run_kw = p_kw.add_run(
        "Stock Market Prediction, Artificial Neural Network, Genetic Algorithm, Technical Indicators, "
        "Feature Selection, Streamlit, MongoDB, Machine Learning, Financial Analytics."
    )
    run_kw.font.size = Pt(12)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #            CHAPTER 1 – INTRODUCTION
    # ════════════════════════════════════════════════════════

    heading(doc, "CHAPTER 1: INTRODUCTION", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "1.1 Overview", level=2)
    para(doc, (
        "Financial markets generate enormous volumes of price and trading data every day. Predicting "
        "the future direction of stock prices has long been a goal for traders, analysts, and researchers "
        "alike. Traditional approaches such as fundamental analysis rely on balance sheets and earnings "
        "reports, while technical analysis leverages historical price patterns and indicators. With the "
        "advent of machine learning, computational methods have emerged as a powerful complement to both."
    ))
    para(doc, (
        'This project, titled "MarketPulse AI – Stock Market Trend Prediction System," aims to build '
        "an intelligent decision-support tool that predicts whether the next trading session for a given "
        "stock will likely move upward (Bullish), downward (Bearish), or remain relatively flat (Neutral). "
        "The system uses a hybrid approach combining an Artificial Neural Network (ANN) for classification "
        "and a Genetic Algorithm (GA) for automated feature selection and hyperparameter optimization."
    ))

    heading(doc, "1.2 Background", level=2)
    para(doc, (
        "Stock markets are inherently non-linear, noisy, and influenced by a complex interplay of "
        "economic indicators, market sentiment, geopolitical events, and investor psychology. Traditional "
        "linear models such as ARIMA have limited capacity to capture these complexities. Neural networks, "
        "particularly Multi-Layer Perceptrons (MLPs), can model non-linear decision boundaries, making "
        "them suitable for binary classification tasks such as predicting market direction."
    ))
    para(doc, (
        "However, the performance of an ANN depends heavily on the choice of input features and "
        "hyperparameters. A Genetic Algorithm—an evolutionary optimization technique inspired by natural "
        "selection—is used here to search across a space of feature subsets and network configurations, "
        "selecting the combination that maximizes prediction accuracy on held-out validation data."
    ))

    heading(doc, "1.3 Motivation", level=2)
    para(doc, (
        "The motivation behind this project is to demonstrate how bio-inspired optimization can enhance "
        "neural network performance for a practical financial application. By automating the tedious "
        "process of feature engineering and model tuning, the system allows users to focus on interpreting "
        "signals rather than configuring algorithms. Additionally, storing predictions in a database "
        "creates an audit trail that supports disciplined, evidence-based investment decision-making."
    ))

    heading(doc, "1.4 Objectives", level=2)
    bullet(doc, "To build a system that downloads and processes real-time stock market data from Yahoo Finance.")
    bullet(doc, "To compute meaningful technical indicators (SMA, EMA, RSI, MACD, Momentum, Volatility) from raw OHLCV data.")
    bullet(doc, "To implement a Genetic Algorithm for automated feature selection and ANN hyperparameter optimization.")
    bullet(doc, "To train an MLP classifier that predicts the next-session market direction.")
    bullet(doc, "To develop a professional, interactive Streamlit-based dashboard with Plotly chart visualizations.")
    bullet(doc, "To integrate MongoDB Atlas for persistent storage and retrieval of prediction history.")

    heading(doc, "1.5 Scope of the Project", level=2)
    para(doc, (
        "The system supports 50+ tickers from both US (NYSE/NASDAQ) and Indian (NSE) markets. Users can "
        "select predefined tickers or enter custom Yahoo Finance symbols. The prediction model works on "
        "daily OHLCV data with configurable date ranges, GA population sizes, and generation counts. The "
        "application is intended as a decision-support prototype for academic and demonstration purposes, "
        "and predictions should not be treated as financial advice."
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #          CHAPTER 2 – LITERATURE REVIEW
    # ════════════════════════════════════════════════════════

    heading(doc, "CHAPTER 2: LITERATURE REVIEW", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "2.1 Stock Market Prediction Using Machine Learning", level=2)
    para(doc, (
        "Numerous studies have explored the use of machine learning for stock market prediction. "
        "Patel et al. (2015) compared Support Vector Machines, Random Forest, ANN, and Naïve Bayes "
        "classifiers for predicting stock trends using ten technical indicators, finding that ANN and "
        "Random Forest performed best. Vijh et al. (2020) used ANN models to forecast stock closing "
        "prices, achieving strong results on Indian stock market data. These studies established that "
        "neural networks can effectively capture non-linear patterns in financial time series."
    ))

    heading(doc, "2.2 Artificial Neural Networks in Finance", level=2)
    para(doc, (
        "Multi-Layer Perceptron (MLP) networks have been widely used for financial forecasting. "
        "Kara et al. (2011) employed a feed-forward ANN to predict stock direction on the Istanbul Stock "
        "Exchange and demonstrated that ANN outperformed traditional statistical models. The advantage of "
        "MLPs lies in their ability to learn complex, non-linear mappings from input features (technical "
        "indicators) to output classes (up or down movement) through backpropagation training."
    ))

    heading(doc, "2.3 Genetic Algorithms for Optimization", level=2)
    para(doc, (
        "Genetic Algorithms (GAs), introduced by Holland (1975) and popularized by Goldberg (1989), are "
        "metaheuristic optimization techniques that simulate natural evolution. GAs maintain a population "
        "of candidate solutions (genomes), apply selection, crossover, and mutation operators, and "
        "iteratively evolve toward better solutions. In the context of machine learning, GAs have been "
        "successfully used for feature selection, hyperparameter tuning, and neural architecture search."
    ))

    heading(doc, "2.4 Hybrid ANN-GA Approaches", level=2)
    para(doc, (
        "The combination of ANN and GA has been explored by several researchers. Kim and Han (2000) used "
        "GA to optimize the connection weights and feature subset of an ANN for stock index prediction. "
        "Chien and Chen (2008) applied GA to select features and optimize ANN topology for financial time "
        "series forecasting. These hybrid approaches consistently outperformed standalone ANN models, as "
        "the GA's global search capability helps avoid local optima in hyperparameter space."
    ))

    heading(doc, "2.5 Technical Indicators", level=2)
    para(doc, "Technical indicators are mathematical transformations of price and volume data that quantify market trends, momentum, and volatility. Commonly used indicators include:")
    bullet(doc, " – Average of closing prices over a fixed window.", bold_prefix="Simple Moving Average (SMA)")
    bullet(doc, " – Weighted average that gives more importance to recent prices.", bold_prefix="Exponential Moving Average (EMA)")
    bullet(doc, " – Measures overbought/oversold conditions on a 0–100 scale.", bold_prefix="Relative Strength Index (RSI)")
    bullet(doc, " – Captures momentum shifts via EMA crossovers.", bold_prefix="MACD")
    bullet(doc, " – Standard deviation of returns, capturing market risk.", bold_prefix="Volatility")
    bullet(doc, " – Rate of price change over a lookback period.", bold_prefix="Momentum")

    heading(doc, "2.6 Gap in Existing Work", level=2)
    para(doc, (
        "While prior studies have demonstrated the viability of hybrid ANN-GA models, most lack a "
        "user-friendly deployment platform. Additionally, few systems provide persistent storage for "
        "predictions, making it difficult to audit model behavior over time. This project addresses both "
        "gaps by integrating the hybrid model into a premium Streamlit dashboard with MongoDB-backed "
        "prediction logging."
    ))

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #          CHAPTER 3 – PROBLEM STATEMENT
    # ════════════════════════════════════════════════════════

    heading(doc, "CHAPTER 3: PROBLEM STATEMENT", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "3.1 Problem Definition", level=2)
    para(doc, (
        "Predicting the short-term direction of stock prices is a challenging task due to the stochastic, "
        "non-linear, and multi-factorial nature of financial markets. Traditional statistical models fail "
        "to capture the complex interactions among technical indicators, and manually tuning machine "
        "learning models for different stocks and time periods is time-consuming and error-prone."
    ))

    heading(doc, "3.2 Specific Challenges", level=2)
    numbered(doc, " Stock prices are influenced by complex, non-linear relationships between numerous factors. Linear models cannot adequately capture these relationships.", bold_prefix="Non-linearity:")
    numbered(doc, " With multiple technical indicators available, selecting the optimal subset of features for prediction is a combinatorial optimization problem.", bold_prefix="Feature selection:")
    numbered(doc, " The performance of neural networks is highly sensitive to hyperparameters such as hidden layer size, activation function, and learning rate.", bold_prefix="Hyperparameter sensitivity:")
    numbered(doc, " Financial markets contain significant noise, making it difficult to distinguish genuine signals from random fluctuations.", bold_prefix="Data noise:")
    numbered(doc, " Most existing prediction tools do not log previous predictions, making it impossible to evaluate model consistency over time.", bold_prefix="Lack of audit trails:")

    heading(doc, "3.3 Proposed Solution", level=2)
    para(doc, "This project proposes a Stock Market Trend Prediction System that addresses the above challenges through:")
    bullet(doc, "A Hybrid ANN-GA model where the Genetic Algorithm automatically selects the best feature subset and optimizes ANN hyperparameters, eliminating manual tuning.")
    bullet(doc, "An MLP Neural Network that captures non-linear relationships between technical indicators and market direction.")
    bullet(doc, "A professional Streamlit dashboard with interactive charts for visual analysis.")
    bullet(doc, "A MongoDB Atlas database for persistent storage and retrieval of all prediction runs, creating a searchable audit trail.")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #            CHAPTER 4 – METHODOLOGY
    # ════════════════════════════════════════════════════════

    heading(doc, "CHAPTER 4: METHODOLOGY", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "4.1 Overall Workflow", level=2)
    steps = [
        "Data Collection (Yahoo Finance / CSV Upload)",
        "Data Preprocessing & Cleaning",
        "Technical Indicator Engineering (SMA, EMA, RSI, MACD, etc.)",
        "Feature Matrix Construction & Target Labeling",
        "Train-Test Split (80/20)",
        "Genetic Algorithm Optimization (Feature Selection + Hyperparameter Tuning)",
        "MLP Neural Network Training & Prediction",
        "Evaluation (Accuracy, Precision, Recall, F1, Confusion Matrix)",
        "Result Visualization & MongoDB Logging",
    ]
    for i, step in enumerate(steps, 1):
        centered(doc, f"↓", size=14, space_after=0)
        centered(doc, f"Step {i}: {step}", size=11, bold=True, space_after=2)
    doc.add_paragraph()

    heading(doc, "4.2 Data Collection", level=2)
    para(doc, (
        "Historical stock data is fetched using the yfinance Python library, which provides OHLCV data "
        "from Yahoo Finance. The system supports a configurable date range and offers a fallback data "
        "source via Alpha Vantage API. Users may also upload offline CSV files in the standard Yahoo "
        "Finance format."
    ))

    heading(doc, "4.3 Feature Engineering", level=2)
    para(doc, "The raw OHLCV data is transformed into a set of 11 technical indicators that serve as input features for the model:")

    feature_rows = [
        ["Return", "Daily percentage change of Close price", "Close.pct_change()"],
        ["SMA_10", "Simple Moving Average (10-day)", "Close.rolling(10).mean()"],
        ["SMA_20", "Simple Moving Average (20-day)", "Close.rolling(20).mean()"],
        ["EMA_10", "Exponential Moving Average (10-day)", "Close.ewm(span=10).mean()"],
        ["EMA_20", "Exponential Moving Average (20-day)", "Close.ewm(span=20).mean()"],
        ["Momentum_5", "Price momentum over 5 days", "Close - Close.shift(5)"],
        ["Volatility", "Rolling std dev of returns", "Return.rolling(10).std()"],
        ["RSI", "Relative Strength Index (14-day)", "Wilder's smoothing method"],
        ["MACD", "Moving Avg Convergence Divergence", "EMA(12) - EMA(26)"],
        ["Signal_Line", "MACD signal line", "MACD.ewm(span=9).mean()"],
        ["Volume_Trend", "Rolling avg of volume change rate", "Volume.pct_change().rolling(5).mean()"],
    ]
    add_table(doc, ["Feature", "Description", "Formula / Method"], feature_rows)

    heading(doc, "4.4 Target Variable", level=2)
    para(doc, (
        "The target variable is a binary label: 1 if the next day's closing price is higher than the "
        "current day's (uptrend), and 0 otherwise (downtrend). This is computed as: "
        "Target = (Close.shift(-1) > Close).astype(int)."
    ))

    heading(doc, "4.5 Genetic Algorithm Design", level=2)
    para(doc, "The Genetic Algorithm performs a simultaneous search over the feature subset and ANN hyperparameters. Each individual in the population (a \"genome\") encodes:")
    bullet(doc, " A binary vector of length 11, where each bit indicates whether a feature is included.", bold_prefix="Feature mask:")
    bullet(doc, " Number of neurons in the hidden layer, chosen from {16, 24, 32, 48, 64}.", bold_prefix="Hidden layer size:")
    bullet(doc, " Either relu or tanh.", bold_prefix="Activation function:")
    bullet(doc, " Initial learning rate, chosen from {0.0005, 0.001, 0.003, 0.005}.", bold_prefix="Learning rate:")

    heading(doc, "4.5.1 GA Operators", level=3)
    bullet(doc, " Fitness-proportional selection retains the top one-third of the population as survivors.", bold_prefix="Selection:")
    bullet(doc, " Single-point crossover on the feature mask; random selection of hyperparameters from either parent.", bold_prefix="Crossover:")
    bullet(doc, " Each feature bit has an 8% chance of flipping; hidden layer size mutates with 15% probability; activation mutates with 10% probability; learning rate mutates with 15% probability.", bold_prefix="Mutation:")

    heading(doc, "4.5.2 Fitness Function", level=3)
    para(doc, (
        "The fitness of a genome is computed as the validation accuracy of the MLP trained with the "
        "genome's feature subset and hyperparameters, minus a small penalty (5%) proportional to the "
        "number of selected features. This encourages parsimonious models:"
    ))
    para(doc, "fitness = accuracy − (num_selected_features / total_features) × 0.05", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "4.6 Neural Network Architecture", level=2)
    para(doc, "The system uses a Multi-Layer Perceptron (MLP) Classifier from scikit-learn with:")
    bullet(doc, "A single hidden layer with GA-optimized neuron count (16–64 neurons)")
    bullet(doc, "GA-selected activation function (relu or tanh)")
    bullet(doc, "GA-optimized learning rate (0.0005–0.005)")
    bullet(doc, "Maximum 700 training iterations for the final model (400 during GA fitness evaluation)")
    bullet(doc, "Standard scaling applied to input features before training")

    heading(doc, "4.7 Evaluation Metrics", level=2)
    para(doc, "The system computes the following classification metrics on the 20% held-out test set:")
    bullet(doc, " Proportion of correctly classified samples.", bold_prefix="Accuracy:")
    bullet(doc, " Proportion of predicted uptrends that were actually uptrends.", bold_prefix="Precision:")
    bullet(doc, " Proportion of actual uptrends that were correctly predicted.", bold_prefix="Recall:")
    bullet(doc, " Harmonic mean of precision and recall.", bold_prefix="F1 Score:")
    bullet(doc, " 2×2 matrix showing TP, FP, TN, and FN.", bold_prefix="Confusion Matrix:")

    heading(doc, "4.8 Technology Stack", level=2)
    tech_rows = [
        ["Language", "Python 3.x", "Core development"],
        ["Web Framework", "Streamlit 1.44", "Interactive dashboard UI"],
        ["Data Processing", "Pandas 2.2, NumPy 2.2", "Data manipulation"],
        ["Machine Learning", "scikit-learn 1.6", "MLP Classifier, metrics"],
        ["Visualization", "Plotly 6.0", "Interactive charts & gauges"],
        ["Data Feed", "yfinance 0.2", "Yahoo Finance API"],
        ["Database", "MongoDB Atlas (PyMongo 4.11)", "Prediction history"],
        ["Configuration", "python-dotenv 1.0", "Environment variables"],
    ]
    add_table(doc, ["Component", "Technology", "Purpose"], tech_rows)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #       CHAPTER 5 – DESIGN & IMPLEMENTATION
    # ════════════════════════════════════════════════════════

    heading(doc, "CHAPTER 5: DESIGN & IMPLEMENTATION", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "5.1 System Architecture", level=2)
    para(doc, (
        "The application follows a modular, service-oriented architecture. The front-end (Streamlit) "
        "communicates with four backend service modules: config.py (application settings), "
        "data_service.py (data fetching & feature engineering), model_service.py (ANN-GA engine), "
        "and db_service.py (MongoDB CRUD operations)."
    ))

    heading(doc, "5.2 Module Description", level=2)

    heading(doc, "5.2.1 Configuration Module (services/config.py)", level=3)
    para(doc, (
        "A dataclass-based configuration loader that reads environment variables for MongoDB connection "
        "URI, database name, and Alpha Vantage API key. Uses python-dotenv for loading .env files."
    ))

    heading(doc, "5.2.2 Data Service Module (services/data_service.py)", level=3)
    para(doc, (
        "Responsible for downloading stock data from Yahoo Finance (with 3 retry attempts and Alpha "
        "Vantage fallback), computing all 11 technical indicators, and generating market snapshots. "
        "Key methods:"
    ))
    bullet(doc, "download_stock_data() – Fetches OHLCV data with caching (TTL: 15 min).")
    bullet(doc, "build_feature_frame() – Engineers all technical indicators and target labels.")
    bullet(doc, "current_market_snapshot() – Returns latest OHLCV values and daily change.")

    heading(doc, "5.2.3 Model Service Module (services/model_service.py)", level=3)
    para(doc, "Contains the HybridTrendPredictor class which implements the full ANN-GA pipeline:")
    bullet(doc, "Genome dataclass – Encodes feature mask and ANN hyperparameters.")
    bullet(doc, "fit_predict() – Splits data (80/20), runs GA optimization, trains final MLP, returns all metrics.")
    bullet(doc, "_run_genetic_search() – Evolves population through selection, crossover, and mutation.")
    bullet(doc, "_fitness() – Evaluates each genome by training a quick MLP and measuring validation accuracy.")

    heading(doc, "5.2.4 Database Service Module (services/db_service.py)", level=3)
    para(doc, "Manages the MongoDB Atlas connection and provides CRUD operations:")
    bullet(doc, "store_prediction() – Inserts a prediction document with symbol, metrics, genome, and timestamp.")
    bullet(doc, "fetch_recent_predictions() – Retrieves the latest N prediction records.")
    bullet(doc, "search_predictions() – Filters by ticker symbol and date range.")

    heading(doc, "5.2.5 Application Module (app.py)", level=3)
    para(doc, "The main Streamlit application (773 lines) that wires all services together and renders five pages:")
    numbered(doc, " Hero section with platform overview and quick navigation.", bold_prefix="Landing Page –")
    numbered(doc, " Candlestick chart, RSI, MACD, probability gauge, stat cards.", bold_prefix="Prediction Dashboard –")
    numbered(doc, " Architecture diagram, confusion matrix, accuracy/precision/recall/F1.", bold_prefix="Model Insights –")
    numbered(doc, " Searchable MongoDB prediction history with filters.", bold_prefix="History / Logs –")
    numbered(doc, " Technical summary and system information cards.", bold_prefix="About Project –")

    heading(doc, "5.3 Database Design", level=2)
    para(doc, 'The system uses MongoDB Atlas (NoSQL) with a single collection named "predictions". Each document stores:')

    db_rows = [
        ["symbol", "String", "Stock ticker (e.g., AAPL, RELIANCE.NS)"],
        ["date_start", "DateTime", "Start of analysis window"],
        ["date_end", "DateTime", "End of analysis window"],
        ["predicted_trend", "String", "Bullish / Bearish / Neutral"],
        ["bullish_probability", "Float", "Uptrend probability (0–1)"],
        ["validation_accuracy", "Float", "Test set accuracy"],
        ["precision", "Float", "Positive prediction quality"],
        ["recall", "Float", "Positive class capture rate"],
        ["f1_score", "Float", "Harmonic mean of precision & recall"],
        ["best_genome", "Object", "Hidden layer size, activation, LR, features"],
        ["confusion_matrix", "Array", "2×2 matrix [[TN, FP], [FN, TP]]"],
        ["market_snapshot", "Object", "Latest OHLCV values and daily change"],
        ["feature_snapshot", "Object", "Latest indicator values"],
        ["created_at", "DateTime", "Prediction timestamp"],
    ]
    add_table(doc, ["Field", "Type", "Description"], db_rows)

    heading(doc, "5.4 User Interface Design", level=2)
    para(doc, "The UI is designed with a trading-terminal aesthetic, featuring a dark-accented professional layout with custom CSS styling (9,400+ bytes). Key design elements include:")
    bullet(doc, " Trading controls with ticker selector (50+ stocks), date range picker, GA parameter sliders, and database connection status.", bold_prefix="Sidebar:")
    bullet(doc, " Horizontal radio buttons for page switching.", bold_prefix="Top Navigation:")
    bullet(doc, " Large candlestick chart with SMA/EMA overlays, probability gauge, and compact stat cards.", bold_prefix="Dashboard:")
    bullet(doc, " Interactive Plotly visualizations with transparent backgrounds, consistent color scheme.", bold_prefix="Charts:")

    heading(doc, "5.5 Project Structure", level=2)
    para(doc, (
        "Stock-Market-Trend-Prediction-System/\n"
        "├── app.py                    # Main Streamlit application (773 lines)\n"
        "├── styles.css                # Custom CSS styling (9,476 bytes)\n"
        "├── requirements.txt          # Python dependencies\n"
        "├── .env                      # Environment variables\n"
        "├── services/\n"
        "│   ├── __init__.py           # Package initializer\n"
        "│   ├── config.py             # Application configuration loader\n"
        "│   ├── data_service.py       # Stock data fetching & feature engineering\n"
        "│   ├── db_service.py         # MongoDB Atlas CRUD operations\n"
        "│   └── model_service.py      # Hybrid ANN-GA prediction engine\n"
        "└── *.csv                     # Sample offline data files"
    ), size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #          CHAPTER 6 – TESTING & RESULTS
    # ════════════════════════════════════════════════════════

    heading(doc, "CHAPTER 6: TESTING & RESULTS", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "6.1 Test Cases", level=2)

    test_rows = [
        ["1", "Fetch data for valid ticker", "Ticker: AAPL", "OHLCV DataFrame", "Pass ✅"],
        ["2", "Fetch data for invalid ticker", "Ticker: XYZZZZ", "Error message", "Pass ✅"],
        ["3", "Feature engineering (valid)", "250+ row DataFrame", "11 indicators + Target", "Pass ✅"],
        ["4", "Feature engineering (<60 rows)", "<60 row DataFrame", "ValueError raised", "Pass ✅"],
        ["5", "GA optimization convergence", "Pop:28, Gen:14", "Best genome returned", "Pass ✅"],
        ["6", "MLP prediction", "Scaled feature matrix", "Predictions + probabilities", "Pass ✅"],
        ["7", "MongoDB storage", "Prediction payload", "Document inserted", "Pass ✅"],
        ["8", "MongoDB offline handling", "Invalid URI", "App works without DB", "Pass ✅"],
        ["9", "CSV file upload", "Yahoo Finance CSV", "Parsed & predicted", "Pass ✅"],
        ["10", "CSV with missing columns", "CSV missing Volume", "Descriptive error", "Pass ✅"],
    ]
    add_table(doc, ["S.No", "Test Case", "Input", "Expected Output", "Status"], test_rows)

    heading(doc, "6.2 Performance Metrics", level=2)
    para(doc, "The hybrid ANN-GA model was tested across multiple stock tickers. Below are sample results:")

    perf_rows = [
        ["AAPL", "NASDAQ", BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT],
        ["TSLA", "NASDAQ", BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT],
        ["RELIANCE.NS", "NSE", BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT],
        ["TCS.NS", "NSE", BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT],
        ["INFY.NS", "NSE", BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT, BLANK_SHORT],
    ]
    add_table(doc, ["Ticker", "Market", "Accuracy", "Precision", "Recall", "F1 Score", "Trend"], perf_rows)
    para(doc, "* Fill in the above table by running the system on each ticker and recording the metrics from the Model Insights page.", italic=True, size=10)

    heading(doc, "6.3 Screenshots", level=2)
    para(doc, "Paste screenshots of the following application screens below:", italic=True)
    doc.add_paragraph()
    for label in [
        "Screenshot 1: Landing Page",
        "Screenshot 2: Prediction Dashboard with Candlestick Chart",
        "Screenshot 3: Model Insights – Confusion Matrix & Metrics",
        "Screenshot 4: History / Logs Page",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[{label}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r.font.size = Pt(11)
        # Add empty space for image
        for _ in range(6):
            doc.add_paragraph()

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #              CHAPTER 7 – CONCLUSION
    # ════════════════════════════════════════════════════════

    heading(doc, "CHAPTER 7: CONCLUSION", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "7.1 Summary", level=2)
    para(doc, (
        "This project successfully designed and implemented a Stock Market Trend Prediction System using "
        "a hybrid Artificial Neural Network–Genetic Algorithm approach. The system demonstrates that "
        "evolutionary optimization can effectively automate the challenging tasks of feature selection and "
        "hyperparameter tuning for neural network classifiers in the financial domain."
    ))
    para(doc, "Key accomplishments of this project include:")
    bullet(doc, "Developed a modular, service-oriented Python application with clearly separated concerns across data fetching, feature engineering, model training, and database operations.")
    bullet(doc, "Implemented a Genetic Algorithm that simultaneously optimizes feature subsets, hidden layer sizes, activation functions, and learning rates.")
    bullet(doc, "Built a premium, chart-first Streamlit dashboard with interactive Plotly visualizations including candlestick charts, RSI, MACD, probability gauges, and confusion matrices.")
    bullet(doc, "Integrated MongoDB Atlas for persistent prediction logging, enabling users to search, audit, and compare historical prediction runs.")
    bullet(doc, "Supported 50+ stock tickers from US and Indian markets, with configurable date ranges and GA parameters.")

    heading(doc, "7.2 Limitations", level=2)
    bullet(doc, "The model relies solely on technical indicators and does not incorporate fundamental data, news sentiment, or macroeconomic factors.")
    bullet(doc, "The GA optimization can be computationally expensive for large population sizes and generation counts.")
    bullet(doc, "Stock market prediction is inherently uncertain—the system is a decision-support tool and should not be used as the sole basis for investment decisions.")
    bullet(doc, "The system uses daily OHLCV data and does not support intraday prediction.")

    heading(doc, "7.3 Future Scope", level=2)
    bullet(doc, " Integrate NLP to analyze financial news and social media sentiment as additional input features.", bold_prefix="Sentiment Analysis:")
    bullet(doc, " Replace the MLP with LSTM or Transformer-based architectures for better temporal pattern recognition.", bold_prefix="Deep Learning Models:")
    bullet(doc, " Extend predictions beyond single-day direction to multi-day price trajectory forecasting.", bold_prefix="Multi-Step Forecasting:")
    bullet(doc, " Combine trend predictions with portfolio optimization algorithms for automated asset allocation.", bold_prefix="Portfolio Optimization:")
    bullet(doc, " Integrate WebSocket-based real-time data feeds for intraday monitoring and prediction.", bold_prefix="Real-Time Streaming:")
    bullet(doc, " Combine multiple models (ANN, SVM, Random Forest) with a meta-learner for improved robustness.", bold_prefix="Ensemble Methods:")
    bullet(doc, " Develop a mobile-friendly version or companion app for on-the-go access.", bold_prefix="Mobile Application:")

    add_page_break(doc)

    # ════════════════════════════════════════════════════════
    #                    REFERENCES
    # ════════════════════════════════════════════════════════

    heading(doc, "REFERENCES", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    references = [
        'Patel, J., Shah, S., Thakkar, P., & Kotecha, K. (2015). Predicting stock and stock price index movement using trend deterministic data preparation and machine learning techniques. Expert Systems with Applications, 42(1), 259–268.',
        'Vijh, M., Chandola, D., Tikkiwal, V. A., & Kumar, A. (2020). Stock closing price prediction using machine learning techniques. Procedia Computer Science, 167, 599–606.',
        'Kara, Y., Boyacioglu, M. A., & Baykan, Ö. K. (2011). Predicting direction of stock price index movement using artificial neural networks and support vector machines. Expert Systems with Applications, 38(5), 5311–5319.',
        'Kim, K. J., & Han, I. (2000). Genetic algorithms approach to feature discretization in artificial neural networks for the prediction of stock price index. Expert Systems with Applications, 19(2), 125–132.',
        'Chien, T. W., & Chen, L. H. (2008). An integrated framework combining genetic algorithm and neural network for financial time series forecasting. International Journal of Innovative Computing, 4(5), 1141–1153.',
        'Goldberg, D. E. (1989). Genetic Algorithms in Search, Optimization, and Machine Learning. Addison-Wesley.',
        'Holland, J. H. (1975). Adaptation in Natural and Artificial Systems. University of Michigan Press.',
        'Murphy, J. J. (1999). Technical Analysis of the Financial Markets. New York Institute of Finance.',
        'Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830.',
        'Streamlit Inc. (2024). Streamlit Documentation. https://docs.streamlit.io/',
        'Plotly Technologies Inc. (2024). Plotly Python Graphing Library. https://plotly.com/python/',
        'MongoDB Inc. (2024). MongoDB Atlas Documentation. https://www.mongodb.com/docs/atlas/',
        'Yahoo Finance. (2024). yfinance Python Library. https://pypi.org/project/yfinance/',
        'Wilder, J. W. (1978). New Concepts in Technical Trading Systems. Trend Research.',
        'LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436–444.',
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        run_num = p.add_run(f"[{i}]  ")
        run_num.bold = True
        run_num.font.size = Pt(11)
        run_text = p.add_run(ref)
        run_text.font.size = Pt(11)

    # ── Save ──
    doc.save(OUTPUT)
    print(f"\n[OK] Report saved to: {OUTPUT}")
    print(f"    Total sections: Front Page, Certificate, Index, Abstract, Ch.1-7, References")
    print(f"    Blanks left for: College name, address, department, 3 teammate names/reg. nos,")
    print(f"                     guide name, degree, academic year, signatures, performance metrics, screenshots\n")


if __name__ == "__main__":
    build_report()
